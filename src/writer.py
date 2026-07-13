from pathlib import Path

from docxtpl import DocxTemplate
from spire.doc import Document, FileFormat, ParagraphStyle, HorizontalAlignment
from docx import Document as Doc

from config import (
    OUTPUT_PATH,
    TEMPORARY_PATH,
    QUESTIONS_PATH,
    ANSWERS_PATH,
    MAIN_COVER_PAGE,
    SUBJECT_COVER_TEMPLATE,
    SUBJECT_COVER_TEMPLATE_ANSWERS,
    TEMPLATE_TITLE_STRING,
    TEMPLATE_ABBREVIATION_STRING,
)
from models import EmployeeData, SubjectData


def create_output_document_path(
    subject: SubjectData,
    employee: EmployeeData,
    is_answers_document: bool = False,
    is_temporary_file: bool = False,
) -> Path:
    """
    Create a file path for a questions or answers document.
    Allows creation of paths in the output folder, or in the temporary folder.
    """
    subject_abbrev = subject["abbreviation"]
    employee_name = employee["name"]
    employee_license = employee["license"]

    if is_temporary_file:
        folder_path = TEMPORARY_PATH
    else:
        folder_path = OUTPUT_PATH / f"{employee_name} {employee_license}"

    folder_path.mkdir(parents=True, exist_ok=True)

    if is_temporary_file:
        file_name = subject_abbrev.lower()
    else:
        file_name = f"{employee_name} {employee_license} {subject_abbrev.upper()}"

    if is_answers_document:
        file_name += " odgovori"

    file_name += ".docx"

    return folder_path / file_name


def create_cover_page_for_subject(
    subject: SubjectData, is_answers_document: bool = False
) -> Path:
    """Create cover page and return temporary file path."""
    if is_answers_document:
        cover_page = DocxTemplate(SUBJECT_COVER_TEMPLATE_ANSWERS)
    else:
        cover_page = DocxTemplate(SUBJECT_COVER_TEMPLATE)

    context = {
        TEMPLATE_TITLE_STRING: subject["title"],
        TEMPLATE_ABBREVIATION_STRING: subject["abbreviation"].upper(),
    }
    cover_page.render(context)

    temp_file = TEMPORARY_PATH / f"cover-{subject['abbreviation']}.docx"
    temp_file.parent.mkdir(parents=True, exist_ok=True)
    cover_page.save(str(temp_file))

    return temp_file


def set_question_number(question_file: str, subject_abbrev: str):
    """
    Set question number to be the same as the question document file name.
    """
    # Get file name from the file path, ignoring folders and extension
    number = question_file.split("/")[-1].split(".")[0]

    # Create file path for temporary question document
    temp_question_path = TEMPORARY_PATH / f"{number}.docx"

    # Create temporary question file
    question_doc = Document()
    question_doc.LoadFromFile(question_file)

    # Get table and cell references
    table = question_doc.Sections[0].Tables[0]
    cell = table.Rows[1].Cells[0]

    # Delete current question number
    cell.Paragraphs.RemoveAt(0)

    # Insert number based on file name
    para = cell.AddParagraph()
    para.AppendText(f"{number}.")

    # Set paragraph style
    style = ParagraphStyle(question_doc)
    style.Name = "Redni broj"
    style.CharacterFormat.FontName = "Arial"
    style.CharacterFormat.FontSize = 18
    style.CharacterFormat.Bold = True
    style.CharacterFormat.Italic = True
    style.ParagraphFormat.HorizontalAlignment = HorizontalAlignment.Center
    question_doc.Styles.Add(style)
    para.ApplyStyle("Redni broj")

    question_doc.SaveToFile(str(temp_question_path), FileFormat.Docx2016)
    question_doc.Close()

    return temp_question_path


def generate_document_for_subject(
    subject: SubjectData,
    employee: EmployeeData,
    is_answers_document: bool = False,
    is_temporary_file: bool = False,
) -> Path:
    """
    Merge cover page with selected question or answer files.
    Save it in the output folder, or in the temporary folder.
    """
    subject_abbrev = subject["abbreviation"]

    # Get output document file path
    if is_temporary_file:
        if is_answers_document:
            output_file_path = create_output_document_path(
                subject, employee, True, True
            )
        else:
            output_file_path = create_output_document_path(
                subject, employee, False, True
            )
    else:
        if is_answers_document:
            output_file_path = create_output_document_path(subject, employee, True)
        else:
            output_file_path = create_output_document_path(subject, employee)

    # Create cover page
    if is_answers_document:
        cover_page_path = create_cover_page_for_subject(subject, True)
    else:
        cover_page_path = create_cover_page_for_subject(subject)

    # Create file names from generated numbers
    question_numbers = subject["generated_numbers"]

    # File names for answers
    if is_answers_document:
        files = [
            f"{ANSWERS_PATH}/{subject_abbrev}/{number}.docx"
            for number in question_numbers
        ]
    # File names for questions
    else:
        files = [
            f"{QUESTIONS_PATH}/{subject_abbrev}/{number}.docx"
            for number in question_numbers
        ]

    # Create subject document (containing either questions or answers)
    subject_document = Document()
    subject_document.LoadFromFile(str(cover_page_path))

    # Merge individual documents (questions or answers)
    for i, file in enumerate(files):
        # Set question number to be same as file name
        temp_question_file = set_question_number(file, subject_abbrev)
        print(temp_question_file)

        # Append a question or an answer document
        subject_document.InsertTextFromFile(str(temp_question_file), FileFormat.Auto)

    # Save file compatible with Word 2016
    subject_document.SaveToFile(str(output_file_path), FileFormat.Docx2016)
    subject_document.Close()

    # Delete watermark from the top of the document generated by Spire.Doc
    subject_document = Doc(str(output_file_path))
    subject_document.paragraphs[0].clear()
    subject_document.save(output_file_path)

    return output_file_path


def generate_documents_for_all_subjects(
    subjects: list[SubjectData],
    employee: EmployeeData,
    is_answers_document: bool = False,
) -> list[Path]:
    """Generate tests or test answers for all subjects and return list of output file paths."""
    output_paths: list[Path] = []

    for subject in subjects:
        if is_answers_document:
            output_path = generate_document_for_subject(subject, employee, True)
        else:
            output_path = generate_document_for_subject(subject, employee)

        output_paths.append(output_path)

    return output_paths


def generate_one_document_for_all_subjects(
    subjects: list[SubjectData],
    employee: EmployeeData,
    is_answers_document: bool = False,
) -> Path:
    """
    Generate one document that contains questions or answers for each subject.
    Return list of output file path.
    """
    # Create output document folder
    output_file_path = OUTPUT_PATH / f"{employee['name']} {employee['license']}"

    # Append file name and extension, and if the document contains answers
    if is_answers_document:
        output_file_path /= f"{employee['name']} {employee['license']} odgovori.docx"
    else:
        output_file_path /= f"{employee['name']} {employee['license']}.docx"

    # Create list of temporary files created for each subject
    generated_subject_documents = []

    for subject in subjects:
        subject_document_path = generate_document_for_subject(
            subject, employee, is_answers_document, True
        )
        generated_subject_documents.append(subject_document_path)

    # Create main subject document with main cover page
    main_document = Document()
    main_document.LoadFromFile(str(MAIN_COVER_PAGE))

    # Append each temporary subject document to main document
    for document in generated_subject_documents:
        main_document.InsertTextFromFile(str(document), FileFormat.Docx2016)

    # Save file compatible with Word 2016
    main_document.SaveToFile(str(output_file_path), FileFormat.Docx2016)
    main_document.Close()

    # Delete watermark from the top of the document generated by Spire.Doc
    main_document = Doc(str(output_file_path))
    main_document.paragraphs[0].clear()
    main_document.save(output_file_path)

    return output_file_path
