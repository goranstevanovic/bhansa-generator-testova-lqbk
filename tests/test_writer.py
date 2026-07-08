"""Tests for writer module."""

from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document
from spire.doc import *
from spire.doc.common import *

from writer import (
    create_output_document_path,
    create_cover_page,
    generate_document_for_subject,
    generate_documents_for_all_subjects,
)

FIXTURES_PATH = Path("tests/fixtures")
SAMPLE_QUESTIONS_PATH = FIXTURES_PATH / "baza" / "pitanja"
SAMPLE_ANSWERS_PATH = FIXTURES_PATH / "baza" / "odgovori"
SAMPLE_OUTPUT_PATH = FIXTURES_PATH / "output"
SAMPLE_TEMPORARY_PATH = FIXTURES_PATH / "tmp"
SAMPLE_COVER_TEMPLATE = FIXTURES_PATH / "baza" / "predlosci" / "template-naslovna.docx"
SAMPLE_COVER_TEMPLATE_ANSWERS = (
    FIXTURES_PATH / "baza" / "predlosci" / "template-naslovna-odgovori.docx"
)
SAMPLE_COVER_PAGE = SAMPLE_COVER_TEMPLATE
SAMPLE_TEMPLATE_TITLE_STRING = "naziv"
SAMPLE_TEMPLATE_ABBREVIATION_STRING = "skracenica"


# Sample employee
@pytest.fixture
def sample_employee():
    return {"name": "Marko Marković", "license": "ATCO.0123"}


# Sample subject
@pytest.fixture
def sample_subject():
    return {
        "abbreviation": "npo",
        "title": "naziv prve oblasti",
        "total_questions": 10,
        "percentage": 40,
        "generated_numbers": [1, 4, 7, 10],
    }


# Sample subject 2
@pytest.fixture
def sample_subject_2():
    return {
        "abbreviation": "ndo",
        "title": "naziv druge oblasti",
        "total_questions": 8,
        "percentage": 50,
        "generated_numbers": [1, 3, 5, 7],
    }


# Sample subject 3
@pytest.fixture
def sample_subject_3():
    return {
        "abbreviation": "nto",
        "title": "naziv treće oblasti",
        "total_questions": 6,
        "percentage": 60,
        "generated_numbers": [2, 3, 4, 5],
    }


# Sample questions
@pytest.fixture
def sample_questions():
    return {
        "pitanje1": "Naziv prve oblasti: Prvo pitanje",
        "pitanje2": "Naziv prve oblasti: Drugo pitanje",
        "pitanje3": "Naziv prve oblasti: Treće pitanje",
        "pitanje4": "Naziv prve oblasti: Četvrto pitanje",
        "pitanje5": "Naziv prve oblasti: Peto pitanje",
        "pitanje6": "Naziv prve oblasti: Šesto pitanje",
        "pitanje7": "Naziv prve oblasti: Sedmo pitanje",
        "pitanje8": "Naziv prve oblasti: Osmo pitanje",
        "pitanje9": "Naziv prve oblasti: Deveto pitanje",
        "pitanje10": "Naziv prve oblasti: Deseto pitanje",
    }


# Sample answers
@pytest.fixture
def sample_answers():
    return {
        "odgovor1": "Naziv prve oblasti: Prvo pitanje - odgovor",
        "odgovor2": "Naziv prve oblasti: Drugo pitanje - odgovor",
        "odgovor3": "Naziv prve oblasti: Treće pitanje - odgovor",
        "odgovor4": "Naziv prve oblasti: Četvrto pitanje - odgovor",
        "odgovor5": "Naziv prve oblasti: Peto pitanje - odgovor",
        "odgovor6": "Naziv prve oblasti: Šesto pitanje - odgovor",
        "odgovor7": "Naziv prve oblasti: Sedmo pitanje - odgovor",
        "odgovor8": "Naziv prve oblasti: Osmo pitanje - odgovor",
        "odgovor9": "Naziv prve oblasti: Deveto pitanje - odgovor",
        "odgovor10": "Naziv prve oblasti: Deseto pitanje - odgovor",
    }


# Tests for create_output_document_path()
class TestCreateOutputDocumentPath:
    @patch("writer.OUTPUT_PATH", SAMPLE_OUTPUT_PATH)
    def test_creates_correct_output_document_path(
        self, sample_employee, sample_subject
    ):
        result = create_output_document_path(sample_subject, sample_employee)

        expected_result = (
            Path(SAMPLE_OUTPUT_PATH)
            / f"{sample_employee['name']} {sample_employee['license']}"
            / f"{sample_employee['name']} {sample_employee['license']} "
            f"{sample_subject['abbreviation'].upper()}.docx"
        )

        assert result == expected_result

    @patch("writer.OUTPUT_PATH", SAMPLE_OUTPUT_PATH)
    def test_creates_folder_if_not_exists(self, sample_subject, sample_employee):
        result = create_output_document_path(sample_subject, sample_employee)

        assert result.parent.exists()


# Tests for create_cover_page()
class TestCreateCoverPage:
    @patch("writer.COVER_TEMPLATE", SAMPLE_COVER_TEMPLATE)
    @patch("writer.TEMPORARY_PATH", SAMPLE_TEMPORARY_PATH)
    @patch("writer.TEMPLATE_TITLE_STRING", SAMPLE_TEMPLATE_TITLE_STRING)
    @patch("writer.TEMPLATE_ABBREVIATION_STRING", SAMPLE_TEMPLATE_ABBREVIATION_STRING)
    def test_creates_cover_page_file(self, sample_subject):
        result = create_cover_page(sample_subject)

        assert result.exists()
        assert result.suffix == ".docx"

    @patch("writer.COVER_PAGE", SAMPLE_COVER_PAGE)
    @patch("writer.COVER_TEMPLATE", SAMPLE_COVER_TEMPLATE)
    @patch("writer.TEMPORARY_PATH", SAMPLE_TEMPORARY_PATH)
    @patch("writer.TEMPLATE_TITLE_STRING", SAMPLE_TEMPLATE_TITLE_STRING)
    @patch("writer.TEMPLATE_ABBREVIATION_STRING", SAMPLE_TEMPLATE_ABBREVIATION_STRING)
    def test_cover_page_for_tests_contains_correct_text(self, sample_subject):
        result = create_cover_page(sample_subject)

        # Open created document
        doc = Document()
        doc.LoadFromFile(str(result))

        # Get all text from document
        full_text = doc.GetText().lower()

        # Assert text is present
        assert "naziv prve oblasti" in full_text
        assert "npo" in full_text
        assert "odgovori" not in full_text

    @patch("writer.COVER_TEMPLATE_ANSWERS", SAMPLE_COVER_TEMPLATE_ANSWERS)
    @patch("writer.TEMPORARY_PATH", SAMPLE_TEMPORARY_PATH)
    @patch("writer.TEMPLATE_TITLE_STRING", SAMPLE_TEMPLATE_TITLE_STRING)
    @patch("writer.TEMPLATE_ABBREVIATION_STRING", SAMPLE_TEMPLATE_ABBREVIATION_STRING)
    def test_cover_page_for_test_answers_contains_correct_text(self, sample_subject):
        result = create_cover_page(sample_subject, True)

        # Open created document
        doc = Document(result)

        # Get all text from document
        full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # Assert text is present
        assert "naziv prve oblasti" in full_text
        assert "npo" in full_text
        assert "odgovori" in full_text


# Tests for generate_document_for_subject()
class TestGenerateDocumentForSubject:
    @patch("writer.OUTPUT_PATH", SAMPLE_OUTPUT_PATH)
    @patch("writer.COVER_TEMPLATE", SAMPLE_COVER_TEMPLATE)
    @patch("writer.COVER_TEMPLATE_ANSWERS", SAMPLE_COVER_TEMPLATE_ANSWERS)
    @patch("writer.TEMPORARY_PATH", SAMPLE_TEMPORARY_PATH)
    @patch("writer.TEMPLATE_TITLE_STRING", SAMPLE_TEMPLATE_TITLE_STRING)
    @patch("writer.TEMPLATE_ABBREVIATION_STRING", SAMPLE_TEMPLATE_ABBREVIATION_STRING)
    @patch("writer.QUESTIONS_PATH", SAMPLE_QUESTIONS_PATH)
    @patch("writer.ANSWERS_PATH", SAMPLE_ANSWERS_PATH)
    def test_generates_test_and_test_answers_files_for_single_subject(
        self, sample_subject, sample_employee
    ):
        result_questions = generate_document_for_subject(
            sample_subject, sample_employee
        )
        result_answers = generate_document_for_subject(
            sample_subject, sample_employee, True
        )

        assert result_questions.exists()
        assert result_questions.suffix == ".docx"
        assert result_answers.exists()
        assert result_questions.suffix == ".docx"

    @patch("writer.OUTPUT_PATH", SAMPLE_OUTPUT_PATH)
    @patch("writer.COVER_TEMPLATE", SAMPLE_COVER_TEMPLATE)
    @patch("writer.TEMPORARY_PATH", SAMPLE_TEMPORARY_PATH)
    @patch("writer.TEMPLATE_TITLE_STRING", SAMPLE_TEMPLATE_TITLE_STRING)
    @patch("writer.TEMPLATE_ABBREVIATION_STRING", SAMPLE_TEMPLATE_ABBREVIATION_STRING)
    @patch("writer.QUESTIONS_PATH", SAMPLE_QUESTIONS_PATH)
    def test_generated_test_file_contains_selected_questions(
        self, sample_subject, sample_employee, sample_questions
    ):
        result = generate_document_for_subject(sample_subject, sample_employee)

        # Open created document
        doc = Document()
        doc.LoadFromFile(str(result))

        # Get all text from documents
        full_text = doc.GetText()

        # Assert selcted questions are in document
        assert sample_questions["pitanje4"] in full_text
        assert sample_questions["pitanje1"] in full_text
        assert sample_questions["pitanje7"] in full_text
        assert sample_questions["pitanje10"] in full_text

        # Assert non-selected questions are not in document
        assert sample_questions["pitanje2"] not in full_text
        assert sample_questions["pitanje3"] not in full_text
        assert sample_questions["pitanje5"] not in full_text
        assert sample_questions["pitanje6"] not in full_text
        assert sample_questions["pitanje8"] not in full_text
        assert sample_questions["pitanje9"] not in full_text

    @patch("writer.OUTPUT_PATH", SAMPLE_OUTPUT_PATH)
    @patch("writer.COVER_TEMPLATE_ANSWERS", SAMPLE_COVER_TEMPLATE_ANSWERS)
    @patch("writer.TEMPORARY_PATH", SAMPLE_TEMPORARY_PATH)
    @patch("writer.TEMPLATE_TITLE_STRING", SAMPLE_TEMPLATE_TITLE_STRING)
    @patch("writer.TEMPLATE_ABBREVIATION_STRING", SAMPLE_TEMPLATE_ABBREVIATION_STRING)
    @patch("writer.ANSWERS_PATH", SAMPLE_ANSWERS_PATH)
    def test_generated_test_answers_file_contains_selected_questions(
        self, sample_subject, sample_employee, sample_answers
    ):
        result = generate_document_for_subject(sample_subject, sample_employee, True)

        # Open created document
        doc = Document()
        doc.LoadFromFile(str(result))

        # Get all text from documents
        full_text = doc.GetText()

        # Assert selcted questions are in document
        assert sample_answers["odgovor4"] in full_text
        assert sample_answers["odgovor1"] in full_text
        assert sample_answers["odgovor7"] in full_text
        assert sample_answers["odgovor10"] in full_text

        # Assert non-selected questions are not in document
        assert sample_answers["odgovor2"] not in full_text
        assert sample_answers["odgovor3"] not in full_text
        assert sample_answers["odgovor5"] not in full_text
        assert sample_answers["odgovor6"] not in full_text
        assert sample_answers["odgovor8"] not in full_text
        assert sample_answers["odgovor9"] not in full_text


# Tests for generate_documents_for_all_subjects()
class TestDocumentsForAllSubjects:
    @patch("writer.OUTPUT_PATH", SAMPLE_OUTPUT_PATH)
    @patch("writer.COVER_TEMPLATE", SAMPLE_COVER_TEMPLATE)
    @patch("writer.TEMPORARY_PATH", SAMPLE_TEMPORARY_PATH)
    @patch("writer.TEMPLATE_TITLE_STRING", SAMPLE_TEMPLATE_TITLE_STRING)
    @patch("writer.TEMPLATE_ABBREVIATION_STRING", SAMPLE_TEMPLATE_ABBREVIATION_STRING)
    @patch("writer.QUESTIONS_PATH", SAMPLE_QUESTIONS_PATH)
    def test_generates_test_files_for_all_subjects(
        self, sample_subject, sample_subject_2, sample_subject_3, sample_employee
    ):
        subjects = [sample_subject, sample_subject_2, sample_subject_3]

        results = generate_documents_for_all_subjects(subjects, sample_employee)

        assert len(results) == 3
        assert all(res.exists() for res in results)

    @patch("writer.OUTPUT_PATH", SAMPLE_OUTPUT_PATH)
    @patch("writer.COVER_TEMPLATE_ANSWERS", SAMPLE_COVER_TEMPLATE_ANSWERS)
    @patch("writer.TEMPORARY_PATH", SAMPLE_TEMPORARY_PATH)
    @patch("writer.TEMPLATE_TITLE_STRING", SAMPLE_TEMPLATE_TITLE_STRING)
    @patch("writer.TEMPLATE_ABBREVIATION_STRING", SAMPLE_TEMPLATE_ABBREVIATION_STRING)
    @patch("writer.ANSWERS_PATH", SAMPLE_ANSWERS_PATH)
    def test_generates_test_answers_files_for_all_subjects(
        self, sample_subject, sample_subject_2, sample_subject_3, sample_employee
    ):
        subjects = [sample_subject, sample_subject_2, sample_subject_3]

        results = generate_documents_for_all_subjects(subjects, sample_employee, True)

        assert len(results) == 3
        assert all(res.exists() for res in results)
